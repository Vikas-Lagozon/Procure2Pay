# tools.py
# ─────────────────────────────────────────────────────────────
# Jarvis — Root Agent Tools
#
# All cross-cutting tool functions used by the root LlmAgent.
# Each function is a self-contained callable that the agent
# can invoke directly. Business logic lives here; the agent
# decides when to call which tool based on user intent.
#
# Tools:
#   list_all_requirements()
#   list_all_vendors()
#   get_requirement_details(record_id)
#   get_vendor_details(record_id)
#   get_requirement_and_all_vendors(requirement_id)
#   get_vendor_and_all_requirements(vendor_id)
#   get_all_requirements_and_all_vendors()
#   save_match_result_to_db(session_id, title, match_data_json)
#   save_match_result_to_docx(title, match_data_json, filename)
# ─────────────────────────────────────────────────────────────

import json
import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import config
from logger import get_logger
from nosql_db import MongoCollection

logger = get_logger(__name__)

# ── MongoDB collections ────────────────────────────────────────
_requirements_col = MongoCollection("requirements")
_vendors_col      = MongoCollection("vendors")
_root_col         = MongoCollection("root")

# ── Storage directory ──────────────────────────────────────────
ROOT_DIR = Path("ROOT")


# ─────────────────────────────────────────────────────────────
# INTERNAL HELPERS
# ─────────────────────────────────────────────────────────────

def _record_summary(rec: dict) -> dict:
    """Slim summary of a DB record — safe to return from a tool."""
    return {
        "record_id":  rec.get("_id", ""),
        "file_name":  rec.get("file_name", ""),
        "created_at": str(rec.get("created_at", "")),
        "updated_at": str(rec.get("updated_at", "")),
        "data":       rec.get("data", {}),
        "raw_text":   rec.get("raw_text", "")[:3000],
    }


def _ensure_root_dir() -> Path:
    """Create ROOT/ directory if it does not exist and return its Path."""
    ROOT_DIR.mkdir(parents=True, exist_ok=True)
    return ROOT_DIR


# ─────────────────────────────────────────────────────────────
# TOOLS
# ─────────────────────────────────────────────────────────────

def list_all_requirements() -> str:
    """
    Fetch a summary list of all requirement records from the database.
    Returns a JSON string with an array of records, each containing
    record_id, file_name, created_at, and structured data.
    Use this to give the user an overview or to look up a record_id
    before performing a match.
    """
    try:
        records   = _requirements_col.fetch_all()
        summaries = [_record_summary(r) for r in records]
        return json.dumps({
            "status":  "ok",
            "count":   len(summaries),
            "records": summaries,
        }, default=str)
    except Exception as e:
        logger.error(f"list_all_requirements failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def list_all_vendors() -> str:
    """
    Fetch a summary list of all vendor records from the database.
    Returns a JSON string with an array of records, each containing
    record_id, file_name, created_at, and structured data.
    Use this to give the user an overview or to look up a record_id
    before performing a match.
    """
    try:
        records   = _vendors_col.fetch_all()
        summaries = [_record_summary(r) for r in records]
        return json.dumps({
            "status":  "ok",
            "count":   len(summaries),
            "records": summaries,
        }, default=str)
    except Exception as e:
        logger.error(f"list_all_vendors failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def get_requirement_details(record_id: str) -> str:
    """
    Fetch the full details of a single requirement record by its
    MongoDB record_id (24-character hex string).
    Returns a JSON string with the complete structured data and raw text.
    """
    try:
        rec = _requirements_col.fetch_by_id(record_id)
        if not rec:
            return json.dumps({"status": "not_found", "record_id": record_id})
        return json.dumps({"status": "ok", "record": _record_summary(rec)}, default=str)
    except Exception as e:
        logger.error(f"get_requirement_details failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def get_vendor_details(record_id: str) -> str:
    """
    Fetch the full details of a single vendor record by its
    MongoDB record_id (24-character hex string).
    Returns a JSON string with the complete structured data and raw text.
    """
    try:
        rec = _vendors_col.fetch_by_id(record_id)
        if not rec:
            return json.dumps({"status": "not_found", "record_id": record_id})
        return json.dumps({"status": "ok", "record": _record_summary(rec)}, default=str)
    except Exception as e:
        logger.error(f"get_vendor_details failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def get_requirement_and_all_vendors(requirement_id: str) -> str:
    """
    Fetch a specific requirement AND all vendor documents in one call.
    Use this when the user asks to find, match, or rank vendors for a
    specific requirement. After receiving the data, analyse and rank the
    vendors by relevance, price, certifications, and technical fit.

    Parameters
    ----------
    requirement_id : str
        MongoDB record_id (24-character hex string) of the requirement.

    Returns
    -------
    str
        JSON string containing the full requirement record and a list
        of all vendor records (with their structured data and raw text).
    """
    try:
        req = _requirements_col.fetch_by_id(requirement_id)
        if not req:
            return json.dumps({"status": "not_found", "requirement_id": requirement_id})

        vendors = _vendors_col.fetch_all()
        return json.dumps({
            "status":       "ok",
            "requirement":  _record_summary(req),
            "vendors":      [_record_summary(v) for v in vendors],
            "vendor_count": len(vendors),
        }, default=str)
    except Exception as e:
        logger.error(f"get_requirement_and_all_vendors failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def get_vendor_and_all_requirements(vendor_id: str) -> str:
    """
    Fetch a specific vendor AND all requirement documents in one call.
    Use this when the user asks which requirements a particular vendor
    can fulfill. After receiving the data, analyse the vendor's product
    categories, technical specs, and pricing to determine which
    requirements it can meet.

    Parameters
    ----------
    vendor_id : str
        MongoDB record_id (24-character hex string) of the vendor.

    Returns
    -------
    str
        JSON string containing the full vendor record and a list of all
        requirement records (with their structured data and raw text).
    """
    try:
        vendor = _vendors_col.fetch_by_id(vendor_id)
        if not vendor:
            return json.dumps({"status": "not_found", "vendor_id": vendor_id})

        requirements = _requirements_col.fetch_all()
        return json.dumps({
            "status":            "ok",
            "vendor":            _record_summary(vendor),
            "requirements":      [_record_summary(r) for r in requirements],
            "requirement_count": len(requirements),
        }, default=str)
    except Exception as e:
        logger.error(f"get_vendor_and_all_requirements failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def get_all_requirements_and_all_vendors() -> str:
    """
    Fetch ALL requirement records and ALL vendor records in one call.
    Use this when the user asks for a comprehensive match table across
    all requirements (e.g., 'top 3 vendors for every requirement').
    After receiving the data, analyse each requirement against all
    vendors and produce the match table.

    Returns
    -------
    str
        JSON string with two lists: requirements and vendors,
        each containing the full structured data and raw text.
    """
    try:
        requirements = _requirements_col.fetch_all()
        vendors      = _vendors_col.fetch_all()
        return json.dumps({
            "status":            "ok",
            "requirements":      [_record_summary(r) for r in requirements],
            "requirement_count": len(requirements),
            "vendors":           [_record_summary(v) for v in vendors],
            "vendor_count":      len(vendors),
        }, default=str)
    except Exception as e:
        logger.error(f"get_all_requirements_and_all_vendors failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def save_match_result_to_db(session_id: str, title: str, match_data_json: str) -> str:
    """
    Save a vendor-requirement match result to the MongoDB 'root'
    collection, keyed by session_id.

    Call this when the user says:
      "Save this result", "Save this table", "Store this match"

    Parameters
    ----------
    session_id : str
        The current session ID — used as the primary key for the record.
    title : str
        A short descriptive title for this match result,
        e.g. "Top 5 vendors for laptop requirement".
    match_data_json : str
        The match result as a JSON string. Structure the data so it
        captures the requirement, the ranked vendors, and the match
        scores or reasoning for each vendor.

    Returns
    -------
    str
        JSON string confirming the save, with the inserted record_id.
    """
    try:
        match_data = json.loads(match_data_json) if isinstance(match_data_json, str) else match_data_json
    except Exception:
        match_data = {"raw": match_data_json}

    record = {
        "session_id": session_id,
        "title":      title,
        "data":       match_data,
        "saved_at":   datetime.datetime.utcnow(),
    }

    try:
        existing = _root_col.fetch_one({"session_id": session_id, "title": title})
        if existing:
            _root_col.update_one(
                {"session_id": session_id, "title": title},
                {"data": match_data, "saved_at": datetime.datetime.utcnow()},
            )
            logger.info(f"Match result updated in DB | session_id={session_id} | title={title}")
            return json.dumps({
                "status":     "updated",
                "session_id": session_id,
                "title":      title,
                "record_id":  existing.get("_id"),
            }, default=str)
        else:
            inserted_id = _root_col.insert_one(record)
            logger.info(f"Match result saved to DB | session_id={session_id} | id={inserted_id}")
            return json.dumps({
                "status":     "saved",
                "session_id": session_id,
                "title":      title,
                "record_id":  str(inserted_id),
            }, default=str)
    except Exception as e:
        logger.error(f"save_match_result_to_db failed: {e}")
        return json.dumps({"status": "error", "message": str(e)})


def save_match_result_to_docx(title: str, match_data_json: str, filename: str) -> str:
    """
    Create a formatted Word document (.docx) in the ROOT/ directory
    containing the match result as a structured table.

    Call this when the user says:
      "Save to Word", "Export to docx", "Save this table to a file",
      "Save this requirement and vendor match table"

    Parameters
    ----------
    title : str
        Document title, printed at the top of the Word file.
    match_data_json : str
        Match result as a JSON string. Expected structure:
        {
          "requirement": { "name": "...", "description": "...", ... },
          "ranked_vendors": [
            {
              "rank": 1,
              "vendor_name": "...",
              "match_score": "...",
              "reason": "...",
              "unit_price": "...",
              "certifications": "...",
              "delivery_terms": "..."
            },
            ...
          ]
        }
        OR for a multi-requirement table:
        {
          "rows": [
            {
              "requirement": "...",
              "vendor_1": "...",
              "vendor_2": "...",
              "vendor_3": "..."
            },
            ...
          ]
        }
    filename : str
        Output filename without extension, e.g. "laptop_vendor_match".
        The file will be saved as ROOT/<filename>.docx.

    Returns
    -------
    str
        JSON string confirming the save with the full file path.
    """
    try:
        data = json.loads(match_data_json) if isinstance(match_data_json, str) else match_data_json
    except Exception:
        data = {"raw": match_data_json}

    _ensure_root_dir()

    # Sanitise filename
    safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in filename).strip()
    if not safe_name:
        safe_name = f"match_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"
    out_path = ROOT_DIR / f"{safe_name}.docx"

    doc = DocxDocument()

    # ── Title ──
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Timestamp ──
    ts = doc.add_paragraph(
        f"Generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}"
    )
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── TABLE: single requirement → ranked vendors ──
    if "ranked_vendors" in data:
        req_info = data.get("requirement", {})
        req_name = req_info.get("name") or req_info.get("item_name") or "Requirement"

        doc.add_heading(f"Requirement: {req_name}", level=2)
        for key, val in req_info.items():
            if key != "name":
                doc.add_paragraph(f"  {key.replace('_', ' ').title()}: {val}")
        doc.add_paragraph("")

        vendors = data.get("ranked_vendors", [])
        if vendors:
            cols  = ["Rank", "Vendor", "Match Score", "Unit Price",
                     "Certifications", "Delivery Terms", "Reason"]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for i, col in enumerate(cols):
                hdr[i].text = col
                for para in hdr[i].paragraphs:
                    for run in para.runs:
                        run.bold = True

            for v in vendors:
                row = table.add_row().cells
                row[0].text = str(v.get("rank", ""))
                row[1].text = str(v.get("vendor_name", ""))
                row[2].text = str(v.get("match_score", ""))
                row[3].text = str(v.get("unit_price", ""))
                row[4].text = str(v.get("certifications", ""))
                row[5].text = str(v.get("delivery_terms", ""))
                row[6].text = str(v.get("reason", ""))

    # ── TABLE: multi-requirement match table ──
    elif "rows" in data:
        rows_data = data.get("rows", [])
        if rows_data:
            cols  = list(rows_data[0].keys())
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for i, col in enumerate(cols):
                hdr[i].text = col.replace("_", " ").title()
                for para in hdr[i].paragraphs:
                    for run in para.runs:
                        run.bold = True

            for row_data in rows_data:
                row = table.add_row().cells
                for i, col in enumerate(cols):
                    row[i].text = str(row_data.get(col, ""))

    # ── Fallback: raw JSON dump ──
    else:
        doc.add_paragraph("Match Data:")
        doc.add_paragraph(json.dumps(data, indent=2, default=str))

    doc.save(str(out_path))
    logger.info(f"Match result saved to docx: {out_path}")

    return json.dumps({
        "status":    "saved",
        "file_path": str(out_path),
        "title":     title,
    }, default=str)


# ─────────────────────────────────────────────────────────────
# TOOL REGISTRY
# Imported by chatbot.py and passed directly to LlmAgent(tools=TOOLS)
# ─────────────────────────────────────────────────────────────

TOOLS = [
    list_all_requirements,
    list_all_vendors,
    get_requirement_details,
    get_vendor_details,
    get_requirement_and_all_vendors,
    get_vendor_and_all_requirements,
    get_all_requirements_and_all_vendors,
    save_match_result_to_db,
    save_match_result_to_docx,
]
